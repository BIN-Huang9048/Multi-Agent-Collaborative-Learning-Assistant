"""
============================================================
FastAPI 应用入口 (main.py)
============================================================
职责：
  1. 创建 FastAPI 应用实例，配置 CORS 跨域中间件
  2. 初始化全局统一日志系统（控制台 + 文件双输出）
  3. 提供健康检查接口 GET /api/health
  4. 文件上传接口 POST /api/upload_file
  5. 笔记生成接口 POST /api/generate_note
  6. 全局请求日志 & 异常处理

架构原则：
  - Agent 负责决策、自检、重试；LLM 仅作为文本生成工具
  - 所有配置从 app.config 模块统一读取
  - 日志全局统一：一个格式、一个级别、两个出口（控制台 + 滚动文件）
  - 业务模块（FileHandler / VectorStore / AgentWorkflow）单例复用
============================================================
"""

import asyncio
import io
import json
import logging
import re as _re
import sys
import threading
import time
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Any, Dict, Optional
from urllib.parse import quote

from fastapi import FastAPI, File, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response, StreamingResponse

# ----------------------------------------------------------
# 导入全局配置（必须在其他业务模块之前，确保 .env 已加载）
# ----------------------------------------------------------
from app.config import (
    # ---- 服务配置 ----
    SERVER_HOST,
    SERVER_PORT,
    DEBUG_MODE,
    # ---- CORS 配置 ----
    CORS_ALLOW_ORIGINS,
    CORS_ALLOW_CREDENTIALS,
    CORS_ALLOW_METHODS,
    CORS_ALLOW_HEADERS,
    # ---- 日志配置 ----
    LOG_DIR,
    LOG_LEVEL,
    LOG_FORMAT,
    LOG_DATE_FORMAT,
    LOG_MAX_BYTES,
    LOG_BACKUP_COUNT,
    # ---- 项目路径 ----
    KNOWLEDGE_DIR,
)

# ----------------------------------------------------------
# 导入笔记历史存储模块
# ----------------------------------------------------------
from app import notes_store

# ----------------------------------------------------------
# 导入数据模型
# ----------------------------------------------------------
from app.schemas import (
    DownloadNoteRequest,
    GenerateNoteRequest,
    GenerateNoteResponse,
    UploadFileResponse,
    ErrorResponse,
)

# ============================================================
# 一、全局日志系统初始化
#    原则：在所有业务代码执行前完成日志配置，确保全局统一输出。
# ============================================================


def setup_logging() -> None:
    """
    配置全局统一日志系统。

    特性：
      - 同时输出到控制台（stdout）和滚动日志文件
      - 日志文件自动轮转（默认单文件 10MB，保留 5 个备份）
      - 统一格式：时间 | 级别 | 模块:函数:行号 | 消息
      - 对第三方库（chromadb、langchain 等）的冗余日志做降噪处理
    """
    # --- 获取根 logger ---
    root_logger = logging.getLogger()
    root_logger.setLevel(getattr(logging, LOG_LEVEL.upper(), logging.INFO))

    # --- 清除已有 handler，避免重复添加（热重载场景） ---
    root_logger.handlers.clear()

    # --- 日志格式器 ---
    formatter = logging.Formatter(
        fmt=LOG_FORMAT,
        datefmt=LOG_DATE_FORMAT,
    )

    # --- 控制台 handler（stdout，开发调试用） ---
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.DEBUG if DEBUG_MODE else logging.INFO)
    console_handler.setFormatter(formatter)
    root_logger.addHandler(console_handler)

    # --- 文件 handler（滚动日志，生产追溯用） ---
    try:
        LOG_DIR.mkdir(parents=True, exist_ok=True)
        log_file_path = LOG_DIR / "app.log"

        file_handler = RotatingFileHandler(
            filename=str(log_file_path),
            maxBytes=LOG_MAX_BYTES,
            backupCount=LOG_BACKUP_COUNT,
            encoding="utf-8",
        )
        file_handler.setLevel(logging.DEBUG)  # 文件记录全量日志，不受 DEBUG_MODE 影响
        file_handler.setFormatter(formatter)
        root_logger.addHandler(file_handler)
    except (OSError, PermissionError) as e:
        # 文件日志创建失败不阻断服务启动，仅输出警告
        print(f"[WARN] 无法创建日志文件 handler: {e}", file=sys.stderr)

    # --- 第三方库日志降噪 ---
    # ChromaDB / LangChain 等库在 DEBUG 模式下会输出大量内部日志，
    # 此处将其级别提升至 WARNING，保持应用日志清晰可读。
    for noisy_lib in (
        "chromadb",
        "chromadb.telemetry",
        "langchain",
        "langgraph",
        "httpx",
        "httpcore",
        "urllib3",
        "asyncio",
    ):
        logging.getLogger(noisy_lib).setLevel(logging.WARNING)

    # --- 确认日志系统就绪 ---
    app_logger = logging.getLogger(__name__)
    app_logger.info("=" * 60)
    app_logger.info(
        "日志系统初始化完成 | 级别=%s | 文件=%s",
        LOG_LEVEL,
        LOG_DIR / "app.log",
    )
    app_logger.info("=" * 60)


# 在模块加载时执行日志初始化
setup_logging()

# 获取本模块对应的 logger
logger = logging.getLogger(__name__)


# ============================================================
# 工具函数
# ============================================================

def _format_size(bytes_val: int) -> str:
    """将字节数格式化为人类可读的大小字符串。"""
    if bytes_val < 1024:
        return f"{bytes_val} B"
    elif bytes_val < 1024 * 1024:
        return f"{bytes_val / 1024:.1f} KB"
    else:
        return f"{bytes_val / (1024 * 1024):.1f} MB"

# ============================================================
# 二、FastAPI 应用实例创建
# ============================================================

app = FastAPI(
    title="基于职能型多Agent的本地知识库学习笔记生成助手",
    description=(
        "支持 TXT / PDF / MD 文件上传，ChromaDB 向量库增量更新，"
        "三大职能 Agent（检索→理解→生成）串行协作生成结构化学习笔记。"
    ),
    version="0.1.0",
    docs_url="/docs" if DEBUG_MODE else None,       # 生产环境可关闭 Swagger
    redoc_url="/redoc" if DEBUG_MODE else None,     # 生产环境可关闭 ReDoc
    debug=DEBUG_MODE,
)

logger.info(
    "FastAPI 应用实例创建完成 | title=%s | debug=%s",
    app.title,
    DEBUG_MODE,
)

# ============================================================
# 三、CORS 跨域中间件配置
# ============================================================

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ALLOW_ORIGINS,
    allow_credentials=CORS_ALLOW_CREDENTIALS,
    allow_methods=CORS_ALLOW_METHODS,
    allow_headers=CORS_ALLOW_HEADERS,
)

logger.info(
    "CORS 中间件已注册 | origins=%s | credentials=%s",
    CORS_ALLOW_ORIGINS,
    CORS_ALLOW_CREDENTIALS,
)

# ============================================================
# 四、请求日志中间件（记录每次请求的方法、路径、耗时）
# ============================================================


@app.middleware("http")
async def request_logging_middleware(request: Request, call_next):
    """
    全局请求日志中间件 —— 记录每次 HTTP 请求的方法、路径、状态码、耗时。

    执行时机：CORS 之后、路由处理之前
    不影响请求/响应内容，纯观测性功能。
    """
    start_time = time.time()

    # 记录请求到达
    logger.debug(
        "--> %s %s | client=%s",
        request.method,
        request.url.path,
        request.client.host if request.client else "unknown",
    )

    # 执行后续处理链（路由 → 业务逻辑）
    response = await call_next(request)

    # 记录响应
    duration = (time.time() - start_time) * 1000  # 毫秒
    logger.info(
        "<-- %s %s | status=%d | %.2fms",
        request.method,
        request.url.path,
        response.status_code,
        duration,
    )

    return response


# ============================================================
# 五、全局异常处理
# ============================================================


@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception) -> JSONResponse:
    """
    捕获所有未处理的异常，统一返回 500 错误 + ErrorResponse 格式。

    原则：
      - 不向客户端泄露敏感堆栈信息
      - DEBUG 模式下附带异常类型名辅助调试
      - 生产模式下仅返回通用 InternalError
    """
    logger.exception(
        "未捕获异常 | path=%s | method=%s | error=%s",
        request.url.path,
        request.method,
        exc,
    )
    error_resp = ErrorResponse(
        code=500,
        msg="服务器内部错误，请稍后重试。",
        detail=str(exc) if DEBUG_MODE else None,
        error_type=type(exc).__name__ if DEBUG_MODE else "InternalError",
    )
    return JSONResponse(
        status_code=500,
        content=error_resp.model_dump(),
    )


# ============================================================
# 六、健康检查接口（保持不变）
# ============================================================


@app.get(
    "/api/health",
    summary="健康检查",
    description="返回服务运行状态，用于监控和负载均衡探活。",
    tags=["系统"],
)
async def health_check() -> Dict[str, str]:
    """
    GET /api/health

    返回格式：
      { "status": "ok", "version": "0.1.0" }
    """
    logger.debug("健康检查请求")
    return {
        "status": "ok",
        "version": app.version,
    }

# ============================================================
# ====================== 以下为新增代码 ======================
# ============================================================
# 七、服务层单例管理
#    复用已有的 FileHandler、VectorStoreManager、Agent 实例，
#    避免每次请求都重新初始化。
# ============================================================


# 模块级单例引用（延迟初始化，首次使用时创建）
_file_handler: Optional["FileHandler"] = None       # type: ignore[name-defined]
_vector_store: Optional["VectorStoreManager"] = None # type: ignore[name-defined]
_services_initialized: bool = False


def _get_file_handler() -> "FileHandler":
    """
    获取 FileHandler 单例。

    入参：无
    出参：FileHandler 实例
    """
    global _file_handler
    if _file_handler is None:
        # 延迟导入，避免模块加载时的循环依赖
        from app.file_handler import FileHandler

        _file_handler = FileHandler()
        logger.info("FileHandler 单例已创建")
    return _file_handler


def _get_vector_store() -> "VectorStoreManager":
    """
    获取 VectorStoreManager 单例。

    出参：VectorStoreManager 实例

    说明：
      - 使用 HuggingFace 本地 Embedding 后端（BAAI/bge-small-zh-v1.5）
      - DeepSeek 不提供 Embedding API，因此使用离线本地模型
      - 若本地模型不可用，自动降级为哑后端（仅测试用）
    """
    global _vector_store
    if _vector_store is None:
        from app.vector_store import VectorStoreManager

        # 尝试 HuggingFace 本地 Embedding 后端；失败则降级
        try:
            _vector_store = VectorStoreManager(embedding_backend="huggingface")
            logger.info("VectorStoreManager 单例已创建 | 后端=HuggingFace 本地模型")
        except (ImportError, RuntimeError) as exc:
            logger.warning(
                "HuggingFace Embedding 后端不可用（%s），降级为哑后端。"
                "向量检索将无语义区分能力，仅用于功能演示。",
                exc,
            )
            _vector_store = VectorStoreManager(embedding_backend="dummy")
            logger.info("VectorStoreManager 单例已创建 | 后端=Dummy")

    return _vector_store


def _ensure_services() -> None:
    """
    确保所有服务层单例已初始化。

    在应用启动时调用一次，预热所有模块。
    失败不阻断启动（各接口内部有自己的错误处理）。
    """
    global _services_initialized
    if _services_initialized:
        return

    try:
        _get_file_handler()
        logger.info("FileHandler 预热完成")
    except Exception as exc:
        logger.warning("FileHandler 预热失败：%s", exc)

    try:
        _get_vector_store()
        logger.info("VectorStoreManager 预热完成")
    except Exception as exc:
        logger.warning("VectorStoreManager 预热失败：%s", exc)

    _services_initialized = True
    logger.info("服务层初始化完成（已就绪=%s）", _services_initialized)


# ============================================================
# 八、业务接口 1：文件上传 + 向量库入库
# ============================================================


@app.post(
    "/api/upload_file",
    summary="文件上传",
    description="上传 TXT / PDF / MD 文件，自动解析、清洗、分片并增量入库到向量库。",
    tags=["知识库"],
    response_model=UploadFileResponse,
    responses={
        200: {"description": "上传并入库成功"},
        400: {"model": ErrorResponse, "description": "文件校验失败"},
        500: {"model": ErrorResponse, "description": "服务器处理异常"},
    },
)
async def api_upload_file(
    file: UploadFile = File(..., description="上传的文件（支持 .txt .pdf .md）"),
) -> UploadFileResponse:
    """
    POST /api/upload_file

    完整流程：
      1. 文件二次校验（格式、大小、非空）
      2. 保存到 knowledge/ 目录
      3. 提取纯文本（TXT → 直接读 / MD → 去格式 / PDF → PyPDF2 逐页提取）
      4. 文本清洗（去乱码、去空行、去多余空白）
      5. 文本分片（LangChain RecursiveCharacterTextSplitter，1000字符/200重叠）
      6. 向量库增量入库（ChromaDB 追加，不重建索引）

    入参：
      file : UploadFile — 上传的文件对象（multipart/form-data）

    返回：
      UploadFileResponse — {
        "code": 200,
        "msg": "文件上传并入库成功",
        "filename": "清洗后文件名",
        "chunks_count": 12
      }
    """
    logger.info("===== [接口] POST /api/upload_file =====")
    logger.info(
        "收到上传请求 | 文件名=%s | content_type=%s",
        file.filename,
        file.content_type,
    )

    # ---- Step 1：文件二次校验 ----
    handler = _get_file_handler()
    try:
        handler.check_file(file, file.filename or "unknown")
    except (ValueError, IOError) as exc:
        logger.warning("文件校验失败 | 文件名=%s | 原因=%s", file.filename, exc)
        return JSONResponse(
            status_code=400,
            content=ErrorResponse(
                code=400,
                msg=str(exc),
                detail=None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )

    # ---- Step 2：保存文件到 knowledge/ ----
    try:
        saved_path = handler.save_file(file, KNOWLEDGE_DIR)
    except (IOError, ValueError) as exc:
        logger.exception("文件保存失败 | 文件名=%s", file.filename)
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"文件保存失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )

    # ---- Step 3：提取纯文本 ----
    try:
        raw_text = handler.extract_text(saved_path)
        logger.info("文本提取完成 | 文件名=%s | 原始长度=%d", saved_path.name, len(raw_text))
    except (IOError, ValueError, FileNotFoundError) as exc:
        logger.exception("文本提取失败 | 文件=%s", saved_path.name)
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"文本提取失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )

    # ---- Step 4：文本清洗 ----
    try:
        clean_text = handler.clean_text(raw_text)
        logger.info("文本清洗完成 | 清洗后长度=%d", len(clean_text))
    except Exception as exc:
        logger.exception("文本清洗失败")
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"文本清洗失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )

    # ---- Step 5：文本分片 ----
    try:
        chunks = handler.split_text(clean_text)
        logger.info("文本分片完成 | 分片数=%d", len(chunks))
    except (ImportError, RuntimeError) as exc:
        logger.exception("文本分片失败")
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"文本分片失败（langchain 未安装？）：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )

    if not chunks:
        logger.warning("分片结果为空 | 文件名=%s", saved_path.name)
        return JSONResponse(
            status_code=400,
            content=ErrorResponse(
                code=400,
                msg="文本分片后无有效内容，请确认文件包含可提取的文本。",
                detail=None,
                error_type="EmptyChunksError",
            ).model_dump(),
        )

    # ---- Step 6：向量库增量入库 ----
    try:
        vector_store = _get_vector_store()
        chunks_count = vector_store.add_documents(
            chunks=chunks,
            source_name=saved_path.name,
        )
        logger.info(
            "向量入库完成 | 文件名=%s | 入库分片=%d",
            saved_path.name,
            chunks_count,
        )
    except (ImportError, RuntimeError) as exc:
        logger.exception("向量库写入失败 | 文件名=%s", saved_path.name)
        # 文件已保存，但向量入库失败 → 返回 500
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"向量库入库失败（chromadb 未安装或写入异常）：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )

    # ---- 成功响应 ----
    logger.info("===== [接口] POST /api/upload_file 完成 =====")
    return UploadFileResponse(
        code=200,
        msg="文件上传并入库成功",
        filename=saved_path.name,
        chunks_count=chunks_count,
    )


# ============================================================
# 九、业务接口 2：笔记生成
# ============================================================


@app.post(
    "/api/generate_note",
    summary="生成学习笔记",
    description="根据用户查询，通过三大 Agent 串行生成结构化学习笔记。",
    tags=["笔记生成"],
    response_model=GenerateNoteResponse,
    responses={
        200: {"description": "笔记生成成功"},
        400: {"model": ErrorResponse, "description": "参数校验失败"},
        500: {"model": ErrorResponse, "description": "服务器处理异常"},
    },
)
async def api_generate_note(
    request: GenerateNoteRequest,
):
    """
    POST /api/generate_note (SSE 流式响应)

    完整流程（由 AgentWorkflow 内部编排）：
      1. RetrievalAgent     — 向量检索 + 相关性自检 + 摘要提纯
      2. UnderstandingAgent — 溯源校验 + 知识框架构建 + 合规自检
      3. NoteGenerateAgent  — 风格模板套用 + 笔记生成 + 格式/内容自检

    通过 Server-Sent Events (SSE) 实时推送每个 Agent 的工作状态。
    最终结果以 event type="result" 推送。

    入参（JSON Body）：
      {
        "query": "什么是RAG？",           // 必填，1~500 字符
        "style": "outline"               // 可选：outline / mindmap / exam_points
      }
    """
    query = request.query.strip()
    style = request.style.strip() or "outline"

    logger.info("===== [接口] POST /api/generate_note (SSE) =====")
    logger.info("收到生成请求 | query=%.80s | style=%s", query, style)

    # ---- Step 1：参数二次校验 ----
    if not query:
        logger.warning("参数校验失败：query 为空")
        return JSONResponse(
            status_code=400,
            content=ErrorResponse(
                code=400,
                msg="查询参数 query 不能为空。",
                detail=None,
                error_type="ValidationError",
            ).model_dump(),
        )

    valid_styles = {"outline", "mindmap", "exam_points"}
    if style not in valid_styles:
        logger.warning("参数校验失败：非法 style=%s", style)
        return JSONResponse(
            status_code=400,
            content=ErrorResponse(
                code=400,
                msg=f"笔记风格不合法「{style}」，可选值：{', '.join(sorted(valid_styles))}",
                detail=None,
                error_type="ValidationError",
            ).model_dump(),
        )

    # ---- Step 2：获取向量库实例 ----
    try:
        vector_store = _get_vector_store()
    except Exception as exc:
        logger.exception("获取 VectorStoreManager 失败")
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"向量库初始化失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )

    # ---- Step 3：SSE 流式执行工作流 ----
    event_queue: asyncio.Queue = asyncio.Queue()
    loop = asyncio.get_running_loop()

    def _status_callback(agent_name: str, phase: str, message: str) -> None:
        """在工作线程中调用，将状态事件安全放入主事件循环队列"""
        loop.call_soon_threadsafe(
            event_queue.put_nowait,
            {"type": "status", "agent": agent_name, "phase": phase, "message": message},
        )

    def _run_workflow_in_thread() -> None:
        """在后台线程中同步执行工作流"""
        try:
            from app.agents import run_workflow

            logger.info("正在调用 Agent 工作流（后台线程）...")
            result = run_workflow(
                query=query,
                style=style,
                vector_store=vector_store,
                status_callback=_status_callback,
            )
            logger.info(
                "工作流返回 | status=%s | note_len=%d | elapsed=%.2fs",
                result.get("status", "unknown"),
                len(result.get("note", "")),
                result.get("elapsed_seconds", 0),
            )

            # 推送最终结果
            loop.call_soon_threadsafe(
                event_queue.put_nowait,
                {
                    "type": "result",
                    "data": {
                        **result,
                        "style": style,   # 补充前端渲染所需的 style 字段
                    },
                },
            )
        except Exception as exc:
            logger.exception("Agent 工作流执行失败（线程内）")
            loop.call_soon_threadsafe(
                event_queue.put_nowait,
                {"type": "error", "message": f"Agent 工作流执行失败：{exc}"},
            )
        finally:
            # 发送哨兵，通知 SSE 生成器结束
            loop.call_soon_threadsafe(event_queue.put_nowait, None)

    # 启动后台线程
    thread = threading.Thread(target=_run_workflow_in_thread, daemon=True)
    thread.start()

    async def _sse_event_generator():
        """SSE 事件生成器 —— 从队列读取事件并序列化为 SSE 格式"""
        while True:
            event: Any = await event_queue.get()

            if event is None:
                # 哨兵：工作流结束
                break

            event_type = event.get("type", "")

            if event_type == "result":
                result_data = event.get("data", {})

                # 自动保存到历史记录
                note_text = result_data.get("note", "")
                if note_text:
                    try:
                        notes_store.save_note(
                            query=query,
                            style=style,
                            context=result_data.get("context", ""),
                            framework=result_data.get("framework", ""),
                            note=note_text,
                        )
                        logger.info("笔记已自动保存到历史 | query=%.60s", query)
                    except Exception as exc:
                        logger.warning("笔记自动保存失败（不影响响应）：%s", exc)

                # 发送最终结果（event: result）
                yield f"event: result\ndata: {json.dumps(result_data, ensure_ascii=False)}\n\n"
                break

            elif event_type == "error":
                yield f"event: error\ndata: {json.dumps(event, ensure_ascii=False)}\n\n"
                break

            elif event_type == "status":
                yield f"event: status\ndata: {json.dumps(event, ensure_ascii=False)}\n\n"

        logger.info("===== [接口] POST /api/generate_note (SSE) 完成 =====")

    return StreamingResponse(
        _sse_event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # 禁用 Nginx 缓冲
        },
    )


# ============================================================
# 十、知识库管理接口（轻量级）
# ============================================================


@app.get(
    "/api/files",
    summary="已入库文件列表",
    description="查询向量库中所有已入库的来源文件名称及物理文件信息。",
    tags=["知识库"],
)
async def api_list_files() -> Dict:
    """
    GET /api/files

    返回：
      {
        "code": 200,
        "msg": "查询成功",
        "files": [
          {
            "filename": "机器学习入门.pdf",
            "size": 1048576,
            "size_display": "1.0 MB",
            "vectors_count": 12
          }
        ],
        "total_vectors": 256
      }
    """
    logger.info("===== [接口] GET /api/files =====")
    try:
        vector_store = _get_vector_store()
        sources = vector_store.get_sources()
        total = vector_store.count()

        # 构造文件详情列表（合并向量库来源 + 物理文件信息）
        files_detail = []
        for source in sources:
            # 统计该来源的向量数
            source_count = 0
            try:
                collection = vector_store._vectorstore._collection
                if collection:
                    result = collection.get(where={"source": source})
                    if result and result.get("ids"):
                        source_count = len(result["ids"])
            except Exception:
                source_count = 0

            # 检查物理文件
            file_path = KNOWLEDGE_DIR / source
            if file_path.exists():
                size_bytes = file_path.stat().st_size
                size_display = _format_size(size_bytes)
            else:
                size_bytes = 0
                size_display = "（仅向量）"

            files_detail.append({
                "filename": source,
                "size": size_bytes,
                "size_display": size_display,
                "vectors_count": source_count,
            })

        logger.info("文件列表查询 | 文件数=%d | 向量总数=%d", len(files_detail), total)
        return {
            "code": 200,
            "msg": "查询成功",
            "files": files_detail,
            "total_vectors": total,
        }
    except Exception as exc:
        logger.exception("文件列表查询失败")
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"查询文件列表失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )


@app.delete(
    "/api/files/{filename}",
    summary="删除指定文件",
    description="从向量库中删除指定文件的所有向量数据，同时删除物理文件。",
    tags=["知识库"],
)
async def api_delete_file(filename: str) -> Dict:
    """
    DELETE /api/files/{filename}

    删除指定文件在向量库中的所有向量数据和本地物理文件。

    返回：
      {
        "code": 200,
        "msg": "文件「xxx.pdf」已删除（12 条向量，物理文件已清理）",
        "filename": "xxx.pdf",
        "deleted_vectors": 12,
        "physical_deleted": true
      }
    """
    logger.info("===== [接口] DELETE /api/files/%s =====", filename)

    # ---- 安全校验：防止路径穿越 ----
    if not filename or ".." in filename or "/" in filename or "\\" in filename:
        return JSONResponse(
            status_code=400,
            content=ErrorResponse(
                code=400,
                msg=f"无效的文件名「{filename}」",
                detail=None,
                error_type="ValidationError",
            ).model_dump(),
        )

    try:
        vector_store = _get_vector_store()

        # 删除向量
        deleted_vectors = vector_store.delete_by_source(filename)

        # 删除物理文件
        handler = _get_file_handler()
        physical_deleted = False
        try:
            physical_deleted = handler.delete_physical_file(filename, KNOWLEDGE_DIR)
        except OSError as exc:
            logger.warning("物理文件删除失败：%s", exc)

        logger.info(
            "文件删除完成 | 文件=%s | 向量删除=%d | 物理删除=%s",
            filename,
            deleted_vectors,
            physical_deleted,
        )

        return {
            "code": 200,
            "msg": f"文件「{filename}」已删除（{deleted_vectors} 条向量"
                   + ("，物理文件已清理）" if physical_deleted else "）"),
            "filename": filename,
            "deleted_vectors": deleted_vectors,
            "physical_deleted": physical_deleted,
        }
    except Exception as exc:
        logger.exception("文件删除失败 | 文件=%s", filename)
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"删除文件失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )


@app.delete(
    "/api/files",
    summary="清空向量库",
    description="删除向量库中的所有数据（谨慎操作，不可恢复）。",
    tags=["知识库"],
)
async def api_clear_files() -> Dict:
    """
    DELETE /api/files

    返回：
      {
        "code": 200,
        "msg": "向量库已清空，共删除 256 条向量。",
        "deleted_count": 256
      }
    """
    logger.info("===== [接口] DELETE /api/files =====")
    try:
        vector_store = _get_vector_store()
        deleted_count = vector_store.clear_vector_store()
        logger.warning("向量库已清空 | 删除数=%d", deleted_count)
        return {
            "code": 200,
            "msg": f"向量库已清空，共删除 {deleted_count} 条向量。",
            "deleted_count": deleted_count,
        }
    except Exception as exc:
        logger.exception("向量库清空失败")
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"向量库清空失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )


# ============================================================
# 十-B、笔记历史管理接口
# ============================================================


@app.get(
    "/api/notes",
    summary="笔记历史列表",
    description="获取已保存的历史笔记摘要列表。",
    tags=["笔记历史"],
)
async def api_list_notes(
    style: str = "",
    limit: int = 50,
    offset: int = 0,
) -> Dict:
    """
    GET /api/notes?style=outline&limit=20&offset=0

    返回：
      {
        "code": 200,
        "msg": "查询成功",
        "total": 10,
        "notes": [
          {
            "id": "abc123",
            "query": "什么是RAG？",
            "style": "outline",
            "created_at": "2025-06-15 14:30:00",
            "note_preview": "一、RAG概述..."
          }
        ]
      }
    """
    logger.info("===== [接口] GET /api/notes | style=%s =====", style or "all")
    try:
        result = notes_store.list_notes(
            style=style if style else None,
            limit=limit,
            offset=offset,
        )
        return {
            "code": 200,
            "msg": "查询成功",
            "total": result["total"],
            "notes": result["notes"],
        }
    except Exception as exc:
        logger.exception("笔记列表查询失败")
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"查询笔记列表失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )


@app.get(
    "/api/notes/{note_id}",
    summary="笔记详情",
    description="获取单条历史笔记的完整内容。",
    tags=["笔记历史"],
)
async def api_get_note(note_id: str) -> Dict:
    """
    GET /api/notes/{note_id}

    返回完整笔记内容（context / framework / note）。
    """
    logger.info("===== [接口] GET /api/notes/%s =====", note_id)
    try:
        note = notes_store.get_note(note_id)
        if note is None:
            return JSONResponse(
                status_code=404,
                content=ErrorResponse(
                    code=404,
                    msg=f"笔记「{note_id}」不存在",
                    detail=None,
                    error_type="NotFoundError",
                ).model_dump(),
            )
        return {
            "code": 200,
            "msg": "查询成功",
            "note": note,
        }
    except Exception as exc:
        logger.exception("笔记详情查询失败 | id=%s", note_id)
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"查询笔记详情失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )


@app.delete(
    "/api/notes/{note_id}",
    summary="删除笔记",
    description="删除指定 ID 的历史笔记。",
    tags=["笔记历史"],
)
async def api_delete_note(note_id: str) -> Dict:
    """
    DELETE /api/notes/{note_id}

    返回：
      {
        "code": 200,
        "msg": "笔记已删除",
        "note_id": "abc123"
      }
    """
    logger.info("===== [接口] DELETE /api/notes/%s =====", note_id)
    try:
        deleted = notes_store.delete_note(note_id)
        if not deleted:
            return JSONResponse(
                status_code=404,
                content=ErrorResponse(
                    code=404,
                    msg=f"笔记「{note_id}」不存在",
                    detail=None,
                    error_type="NotFoundError",
                ).model_dump(),
            )
        return {
            "code": 200,
            "msg": "笔记已删除",
            "note_id": note_id,
        }
    except Exception as exc:
        logger.exception("笔记删除失败 | id=%s", note_id)
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"删除笔记失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )


# ============================================================
# 十-C、笔记下载接口
# ============================================================


@app.post(
    "/api/download_note",
    summary="下载笔记文件",
    description="根据笔记内容和风格生成可下载文件：提纲/考点清单 → .docx，思维导图 → .pdf",
    tags=["笔记下载"],
    responses={
        200: {
            "description": "文件生成成功",
            "content": {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {},
                "application/pdf": {},
            },
        },
        400: {"model": ErrorResponse, "description": "参数校验失败"},
        500: {"model": ErrorResponse, "description": "文件生成失败"},
    },
)
async def api_download_note(request: DownloadNoteRequest):
    """
    POST /api/download_note

    入参（JSON Body）：
      {
        "content": "...",                                 // 笔记内容（必填）
        "style": "outline",                               // outline | mindmap | exam_points
        "svg_markup": "<svg>...</svg>",                   // 思维导图 SVG（仅 mindmap 需要）
        "title": "什么是RAG"                               // 笔记标题
      }

    返回：
      - outline / exam_points → .docx 文件（Word 文档）
      - mindmap               → .pdf 文件（SVG 转换）
    """
    content = request.content
    style = request.style
    svg_markup = request.svg_markup
    # 清理标题，移除可能不安全字符
    title = _re.sub(r'[\\/*?:"<>|]', "_", request.title.strip() or "学习笔记")

    logger.info("===== [接口] POST /api/download_note =====")
    logger.info("下载请求 | style=%s | content_len=%d | title=%.40s", style, len(content), title)

    # ---- Step 1：参数校验 ----
    if not content.strip():
        return JSONResponse(
            status_code=400,
            content=ErrorResponse(
                code=400,
                msg="笔记内容不能为空。",
                detail=None,
                error_type="ValidationError",
            ).model_dump(),
        )

    if style == "mindmap" and not svg_markup.strip():
        return JSONResponse(
            status_code=400,
            content=ErrorResponse(
                code=400,
                msg="思维导图风格需要提供 SVG 标记。",
                detail=None,
                error_type="ValidationError",
            ).model_dump(),
        )

    try:
        if style in ("outline", "exam_points"):
            # ---- DOCX 生成 ----
            file_bytes, mime_type, filename = _generate_docx(content, style, title)
        elif style == "mindmap":
            # ---- PDF 生成（SVG → PDF）----
            file_bytes, mime_type, filename = _generate_pdf_from_svg(
                svg_markup, title
            )
        else:
            return JSONResponse(
                status_code=400,
                content=ErrorResponse(
                    code=400,
                    msg=f"不支持的笔记风格「{style}」",
                    detail=None,
                    error_type="ValidationError",
                ).model_dump(),
            )

        logger.info("下载文件生成成功 | filename=%s | size=%d bytes", filename, len(file_bytes))

        # RFC 5987 编码：HTTP 头仅支持 Latin-1，中文文件名需 URL 编码
        encoded_filename = quote(filename, safe="")
        return Response(
            content=file_bytes,
            media_type=mime_type,
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            },
        )

    except Exception as exc:
        logger.exception("笔记下载文件生成失败 | style=%s | title=%s", style, title)
        return JSONResponse(
            status_code=500,
            content=ErrorResponse(
                code=500,
                msg=f"文件生成失败：{exc}",
                detail=str(exc) if DEBUG_MODE else None,
                error_type=type(exc).__name__,
            ).model_dump(),
        )


# ----------------------------------------------------------
# DOCX 生成（提纲 / 考点清单）
# ----------------------------------------------------------

def _generate_docx(content: str, style: str, title: str):
    """
    将笔记内容生成 .docx 文件。

    入参：
        content : str — 笔记文本（Markdown 格式）
        style   : str — 笔记风格
        title   : str — 文档标题

    出参：
        (bytes, mime_type, filename)
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ---- 标题样式 ----
    title_style = doc.styles["Title"]
    title_font = title_style.font
    title_font.size = Pt(22)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0x4F, 0x1E, 0xA3)

    heading1_style = doc.styles["Heading 1"]
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    heading2_style = doc.styles["Heading 2"]
    heading2_style.font.size = Pt(13)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ---- 添加封面标题 ----
    doc.add_paragraph("")  # 留白
    main_title = doc.add_paragraph(title, style="Title")
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题：风格
    style_names = {"outline": "提纲笔记", "exam_points": "考点清单笔记"}
    subtitle_text = style_names.get(style, "学习笔记")
    subtitle = doc.add_paragraph(subtitle_text)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    # 分隔线
    sep = doc.add_paragraph("─" * 50)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.runs[0]
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph("")  # 间距

    # ---- 解析 Markdown 并添加内容 ----
    _render_markdown_to_docx(doc, content)

    # ---- 页脚 ----
    doc.add_paragraph("")
    footer_line = doc.add_paragraph("— 由本地知识库学习笔记生成助手自动生成 —")
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_line.runs[0]
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer_run.font.italic = True

    # ---- 写入 BytesIO ----
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 文件名（含中文需 URL 编码友好处理，此处直接使用）
    ext_label = "提纲" if style == "outline" else "考点清单"
    filename = f"{title}-{ext_label}.docx"
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return buffer.getvalue(), mime_type, filename


def _render_markdown_to_docx(doc, content: str):
    """
    将 Markdown 文本逐段渲染到 docx 文档中。

    支持：
      - # / ## / ### 标题 → Heading 1/2/3
      - **粗体** / *斜体* 内联标记
      - 【术语】→ 加粗紫色标出
      - 普通段落
    """
    from docx.shared import Pt, RGBColor

    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 标题
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 3")
        elif stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 2")
        elif stripped.startswith("# "):
            p = doc.add_paragraph(stripped[2:], style="Heading 1")
        else:
            # 普通段落
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)

    return doc


def _add_formatted_run(paragraph, text: str):
    """
    给段落添加带格式的 Run，支持：
      - **text** → 粗体
      - *text* → 斜体
      - 【text】 → 紫色加粗
    """
    from docx.shared import Pt, RGBColor

    # 简易状态机解析
    i = 0
    default_size = Pt(10.5)
    while i < len(text):
        # 【术语】
        if text[i:i + 1] == "【":
            end = text.find("】", i)
            if end != -1:
                run = paragraph.add_run(text[i:end + 1])
                run.bold = True
                run.font.size = default_size
                run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
                i = end + 1
                continue

        # **粗体**
        if text[i:i + 2] == "**":
            end = text.find("**", i + 2)
            if end != -1:
                run = paragraph.add_run(text[i + 2:end])
                run.bold = True
                run.font.size = default_size
                i = end + 2
                continue

        # *斜体*
        if text[i:i + 1] == "*" and (i == 0 or text[i - 1] != "*") and text[i:i + 2] != "**":
            end = text.find("*", i + 1)
            if end != -1:
                run = paragraph.add_run(text[i + 1:end])
                run.italic = True
                run.font.size = default_size
                i = end + 1
                continue

        # 普通字符：收集直到下一个特殊标记
        next_special = len(text)
        for marker in ["【", "**", "*"]:
            pos = text.find(marker, i)
            if pos != -1 and pos < next_special:
                next_special = pos
        run = paragraph.add_run(text[i:next_special])
        run.font.size = default_size
        i = next_special

    return paragraph


# ----------------------------------------------------------
# PDF 生成（思维导图 SVG → PDF）
# ----------------------------------------------------------

def _generate_pdf_from_svg(svg_markup: str, title: str):
    """
    将思维导图 SVG 转换为 PDF 文件。

    使用 svglib 解析 SVG → reportlab 图形 → 渲染为 PDF。

    入参：
        svg_markup : str — SVG 标签字符串
        title      : str — 文档标题

    出参：
        (bytes, mime_type, filename)
    """
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPDF
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as pdf_canvas

    # 清理 SVG：去除 HTML 实体和多余空白
    svg_clean = svg_markup.strip()
    # 如果前端传来的是完整的 SVG 元素，直接使用
    if not svg_clean.startswith("<svg") and not svg_clean.startswith("<?xml"):
        # 尝试提取 <svg>...</svg>
        svg_match = _re.search(r"<svg[\s\S]*?</svg>", svg_clean)
        if svg_match:
            svg_clean = svg_match.group(0)

    # 包装为完整的独立 SVG（确保有 xmlns 和 viewBox）
    if 'xmlns="http://www.w3.org/2000/svg"' not in svg_clean and 'xmlns=' not in svg_clean:
        if svg_clean.startswith("<svg"):
            svg_clean = svg_clean[:4] + ' xmlns="http://www.w3.org/2000/svg"' + svg_clean[4:]

    # 如果 SVG 没有 viewBox 但有 width/height，尝试从 viewBox 中获取
    # （前端的 renderMindmapToSvg 总是设置 viewBox）

    # 用 svglib 将 SVG 转为 reportlab 图形对象
    try:
        drawing = svg2rlg(io.BytesIO(svg_clean.encode("utf-8")))
    except Exception as exc:
        logger.warning("svglib 解析 SVG 失败（%s），尝试降级方案", exc)
        # 降级：生成包含 SVG 文字的简单 PDF
        return _generate_fallback_pdf(svg_clean, title)

    if drawing is None:
        logger.warning("svglib 返回 None，降级为简单 PDF")
        return _generate_fallback_pdf(svg_clean, title)

    # 缩放到 A4 页面
    page_w, page_h = A4  # 595.27 x 841.89 points
    margin = 30
    usable_w = page_w - 2 * margin
    usable_h = page_h - 2 * margin

    dw = getattr(drawing, "width", 400) or 400
    dh = getattr(drawing, "height", 300) or 300

    scale = min(usable_w / dw, usable_h / dh, 1.0)
    drawing.width = dw * scale
    drawing.height = dh * scale
    drawing.scale(scale, scale)

    # 居中
    x_offset = margin + (usable_w - drawing.width) / 2
    y_offset = margin + (usable_h - drawing.height) / 2

    # 创建 PDF
    buffer = io.BytesIO()
    c = pdf_canvas.Canvas(buffer, pagesize=A4)
    c.setTitle(title)

    renderPDF.draw(drawing, c, x_offset, page_h - y_offset - drawing.height)

    c.save()
    buffer.seek(0)

    filename = f"{title}-思维导图.pdf"
    return buffer.getvalue(), "application/pdf", filename


def _generate_fallback_pdf(svg_markup: str, title: str) -> tuple:
    """
    降级方案：当 svglib 无法解析 SVG 时，生成包含 Markdown 文本结构的 PDF。
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import cm

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    story = []

    # 标题
    story.append(Paragraph(f"<b>{title}</b> — 思维导图", styles["Title"]))
    story.append(Spacer(1, 0.5 * cm))

    # 尝试从 SVG 中提取文本
    text_matches = _re.findall(r"<text[^>]*>([^<]+)</text>", svg_markup)
    if text_matches:
        story.append(Paragraph("<b>导图节点：</b>", styles["Heading2"]))
        for t in text_matches:
            story.append(Paragraph(f"• {t}", styles["Normal"]))
    else:
        story.append(Paragraph("（无法解析导图结构，请查看在线版思维导图）", styles["Normal"]))

    doc.build(story)
    buffer.seek(0)

    filename = f"{title}-思维导图.pdf"
    return buffer.getvalue(), "application/pdf", filename


# ============================================================
# 十一、应用启动事件
# ============================================================


@app.on_event("startup")
async def on_startup():
    """
    FastAPI 启动事件 —— 预热所有服务层单例。

    在应用开始接受请求前完成：
      - FileHandler 初始化
      - VectorStoreManager 初始化（加载已有向量库）
      - 日志就绪确认
    """
    logger.info("=" * 60)
    logger.info("应用启动中...")
    logger.info("=" * 60)
    _ensure_services()
    logger.info("应用启动完成，等待请求...")


@app.on_event("shutdown")
async def on_shutdown():
    """
    FastAPI 关闭事件 —— 清理资源。

    当前 ChromaDB 自动持久化，无需显式 flush。
    日志轮转由 RotatingFileHandler 自动管理。
    """
    logger.info("应用正在关闭...")
    # ChromaDB 在 VectorStoreManager 析构时自动持久化
    logger.info("应用已关闭。")


# ============================================================
# 十二、启动入口
# ============================================================


def main():
    """
    开发环境启动入口。

    使用方式：
        python -m app.main
    或：
        uvicorn app.main:app --reload --host 0.0.0.0 --port 8000
    """
    import uvicorn

    logger.info("正在启动开发服务器 | host=%s | port=%d", SERVER_HOST, SERVER_PORT)
    uvicorn.run(
        "app.main:app",
        host=SERVER_HOST,
        port=SERVER_PORT,
        reload=DEBUG_MODE,
        log_level=LOG_LEVEL.lower(),
        access_log=DEBUG_MODE,
    )


# 允许直接执行本文件启动服务
if __name__ == "__main__":
    main()
